from docx import Document
import io


# 📝 Generate a structured DOCX meeting summary based on extracted summary data
def generate_docx(summary_data, company_name, meeting_date):
    
    # Create a new Word document instance
    doc = Document()

    # 🏷 Add document title and meeting date
    doc.add_heading(f"{company_name} Meeting Notes", level=0)
    doc.add_paragraph(f"Date: {meeting_date}", style="Heading 2").alignment = 2

    # 1️⃣ Section: Minutes of the Meeting (MoM)
    doc.add_heading("1. Minutes of the Meeting (MoM)", level=1)
    for line in summary_data["mom"]:
        doc.add_paragraph(line.strip(), style="List Bullet")

    # 2️⃣ Section: To-Do List
    doc.add_heading("2. To-Do List", level=1)
    for item in summary_data["todo_list"]:
        doc.add_paragraph(item.strip(), style="List Bullet")

    # 3️⃣ Section: Action Points / Action Plan
    doc.add_heading("3. Action Points / Action Plan", level=1)
    
    # Map keys from the JSON to section titles
    section_titles = {
        "decision_made": "Key Decisions Made",
        "key_services_to_promote": "Key Services to Promote",
        "target_geography": "Target Geography",
        "budget_and_timeline": "Budget and Timeline",
        "lead_management_strategy": "Lead Management Strategy",
        "next_steps_and_ownership": "Next Steps and Ownership",
    }
    
    # ➕ Populate each action plan subsection
    for key, title in section_titles.items():
        doc.add_heading(title, level=2)
        for item in summary_data["action_plan"].get(key, []):
            doc.add_paragraph(item.strip(), style="List Bullet")

    # 💾 Save the document to a binary stream (for uploading/downloading in memory)
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)

    # 📤 Return the document content as a binary object
    return docx_stream.read()