from docx import Document
import io
import re


# 📄 Creates a formatted DOCX document in memory based on a structured summary JSON
def create_docx_in_memory(summary_json, document_title):
    
    # 🧾 Initialize a new Word document
    doc = Document()
    doc.add_heading(document_title, level=0)     # Set the document title as level 0 heading

    # 🔁 Iterate through each structured section in the summary
    for section in summary_json.get("sections", []):
        doc.add_heading(section["heading"], level=1)    # Add each section's heading
        
         # 📌 Process each bullet/content line in the section
        for line in section["content"].split("\n"):
            line = line.strip()
            
            # ➤ Check for markdown-style bullet points
            if line.startswith("- "):
                line = line[2:].strip()     # Remove the "- " prefix
                para = doc.add_paragraph(style="List Bullet")
                parts = re.split(r"(\*\*.*?\*\*)", line)
                
                # ✨ Handle inline bold formatting using **bold**
                for part in parts:
                    run = para.add_run()
                    if part.startswith("**") and part.endswith("**"):
                        run.text = part[2:-2]   # Strip ** and apply bold
                        run.bold = True
                    else:
                        run.text = part
            
            # Add plain text paragraph for non-bullet lines
            else:
                doc.add_paragraph(line.strip())
    
    # 💾 Save the document into an in-memory binary stream
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    # 📤 Return the binary stream (used for uploading or downloading)
    return doc_stream